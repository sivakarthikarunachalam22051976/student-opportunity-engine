
import io
import os
import re
import json
import zipfile
from pathlib import Path
from typing import Any

try:
    from .skill_normalizer import normalize_skills
except ImportError:
    from .skill_normaliser import normalize_skills


# ============================================================
# KNOWN SKILLS DATABASE
# ============================================================

KNOWN_SKILLS = [

    # Programming Languages

    "Python",
    "Java",
    "JavaScript",
    "TypeScript",
    "C++",
    "C",
    "C#",
    "Go",
    "Rust",
    "PHP",
    "Ruby",
    "Kotlin",
    "Swift",
    "R",

    # Artificial Intelligence

    "Machine Learning",
    "Deep Learning",
    "Artificial Intelligence",
    "Generative AI",
    "Large Language Models",
    "LLM",
    "Prompt Engineering",
    "Natural Language Processing",
    "Computer Vision",
    "Neural Networks",

    # Data

    "Data Science",
    "Data Analysis",
    "Data Engineering",
    "Data Visualization",
    "Statistics",
    "Power BI",
    "Tableau",
    "Excel",

    # Databases

    "SQL",
    "MySQL",
    "PostgreSQL",
    "SQLite",
    "MongoDB",
    "Redis",
    "Firebase",

    # Frontend

    "React",
    "React.js",
    "Next.js",
    "Angular",
    "Vue.js",
    "HTML",
    "CSS",
    "Tailwind CSS",
    "Bootstrap",

    # Backend

    "Node.js",
    "Express.js",
    "FastAPI",
    "Django",
    "Flask",
    "Spring Boot",
    "REST APIs",
    "GraphQL",

    # DevOps

    "Docker",
    "Kubernetes",
    "CI/CD",
    "Git",
    "GitHub",
    "GitLab",
    "Linux",
    "Nginx",

    # Cloud

    "AWS",
    "Amazon Web Services",
    "Microsoft Azure",
    "Google Cloud",
    "Google Cloud Platform",

    # AI Frameworks

    "TensorFlow",
    "PyTorch",
    "Scikit-learn",
    "Keras",
    "Hugging Face",
    "LangChain",
    "LlamaIndex",

    # Data Libraries

    "Pandas",
    "NumPy",
    "Matplotlib",
    "Seaborn",

    # APIs and Web

    "API",
    "Web Development",
    "Web Scraping",
    "BeautifulSoup",
    "Selenium",

    # Mobile

    "Android Development",
    "Flutter",
    "React Native",

    # Cybersecurity

    "Cybersecurity",
    "Network Security",
    "Ethical Hacking",

    # Software Engineering

    "Object Oriented Programming",
    "OOP",
    "Data Structures",
    "Algorithms",
    "System Design",
    "Software Development",

    # Tools

    "Jupyter Notebook",
    "Postman",
    "VS Code",
    "Figma",

    # Other Technical Skills

    "Blockchain",
    "Internet of Things",
    "IoT",
    "Cloud Computing",
    "Automation",
]


# ============================================================
# TEXT CLEANING
# ============================================================

def clean_extracted_text(
    text: str,
) -> str:

    if not text:

        return ""


    # Replace null bytes

    text = text.replace(
        "\x00",
        " "
    )


    # Normalize line endings

    text = text.replace(
        "\r\n",
        "\n"
    )

    text = text.replace(
        "\r",
        "\n"
    )


    # Replace tabs

    text = text.replace(
        "\t",
        " "
    )


    # Replace multiple spaces

    text = re.sub(

        r"[ ]{2,}",

        " ",

        text,
    )


    # Replace excessive blank lines

    text = re.sub(

        r"\n{3,}",

        "\n\n",

        text,
    )


    # Remove strange control characters

    text = "".join(

        character

        for character in text

        if character.isprintable()

        or character in "\n\t"
    )


    return text.strip()


# ============================================================
# SAFE TEXT DECODER
# ============================================================

def decode_bytes_safely(
    data: bytes,
) -> str:

    encodings = [

        "utf-8",

        "utf-8-sig",

        "utf-16",

        "utf-16-le",

        "utf-16-be",

        "latin-1",

        "cp1252",
    ]


    for encoding in encodings:

        try:

            text = data.decode(
                encoding
            )


            if text.strip():

                return clean_extracted_text(
                    text
                )


        except Exception:

            continue


    try:

        return clean_extracted_text(

            data.decode(

                "utf-8",

                errors="ignore",
            )
        )

    except Exception:

        return ""


# ============================================================
# PDF EXTRACTION
# ============================================================

def extract_pdf_text(
    file_path: str,
) -> str:

    try:

        from pypdf import PdfReader


        reader = PdfReader(
            file_path
        )


        text_parts = []


        for page in reader.pages:

            try:

                page_text = (
                    page.extract_text()
                    or ""
                )


                if page_text.strip():

                    text_parts.append(
                        page_text
                    )


            except Exception:

                continue


        text = "\n".join(
            text_parts
        )


        return clean_extracted_text(
            text
        )


    except Exception:

        return ""


# ============================================================
# DOCX EXTRACTION
# ============================================================

def extract_docx_text(
    file_path: str,
) -> str:

    try:

        from docx import Document


        document = Document(
            file_path
        )


        text_parts = []


        for paragraph in document.paragraphs:

            if paragraph.text.strip():

                text_parts.append(
                    paragraph.text
                )


        for table in document.tables:

            for row in table.rows:

                row_text = []


                for cell in row.cells:

                    cell_value = (
                        cell.text.strip()
                    )


                    if cell_value:

                        row_text.append(
                            cell_value
                        )


                if row_text:

                    text_parts.append(

                        " | ".join(
                            row_text
                        )
                    )


        return clean_extracted_text(

            "\n".join(
                text_parts
            )
        )


    except Exception:

        return ""


# ============================================================
# OLD DOC EXTRACTION
# ============================================================

def extract_doc_text(
    file_path: str,
) -> str:

    try:

        import textract


        text = textract.process(
            file_path
        )


        return decode_bytes_safely(
            text
        )


    except Exception:

        return ""


# ============================================================
# TXT EXTRACTION
# ============================================================

def extract_txt_text(
    file_path: str,
) -> str:

    try:

        with open(

            file_path,

            "rb",

        ) as file:

            data = file.read()


        return decode_bytes_safely(
            data
        )


    except Exception:

        return ""


# ============================================================
# RTF EXTRACTION
# ============================================================

def extract_rtf_text(
    file_path: str,
) -> str:

    try:

        with open(

            file_path,

            "rb",

        ) as file:

            raw_data = file.read()


        raw_text = decode_bytes_safely(
            raw_data
        )


        # Remove common RTF control words

        text = re.sub(

            r"\\[a-zA-Z]+\d* ?",

            " ",

            raw_text,
        )


        text = text.replace(
            "{",
            " "
        )

        text = text.replace(
            "}",
            " "
        )

        text = text.replace(
            "\\",
            " "
        )


        return clean_extracted_text(
            text
        )


    except Exception:

        return ""


# ============================================================
# HTML EXTRACTION
# ============================================================

def extract_html_text(
    file_path: str,
) -> str:

    try:

        from bs4 import BeautifulSoup


        with open(

            file_path,

            "rb",

        ) as file:

            data = file.read()


        raw_text = decode_bytes_safely(
            data
        )


        soup = BeautifulSoup(

            raw_text,

            "html.parser",
        )


        return clean_extracted_text(

            soup.get_text(

                "\n",

                strip=True,
            )
        )


    except Exception:

        return ""


# ============================================================
# JSON EXTRACTION
# ============================================================

def extract_json_text(
    file_path: str,
) -> str:

    try:

        with open(

            file_path,

            "r",

            encoding="utf-8",

            errors="ignore",

        ) as file:

            data = json.load(
                file
            )


        return clean_extracted_text(

            json.dumps(

                data,

                indent=2,

                ensure_ascii=False,
            )
        )


    except Exception:

        return extract_txt_text(
            file_path
        )


# ============================================================
# GENERIC BINARY / UNKNOWN FILE EXTRACTION
# ============================================================

def extract_generic_text(
    file_path: str,
) -> str:

    try:

        with open(

            file_path,

            "rb",

        ) as file:

            data = file.read()


        # First attempt direct decoding

        text = decode_bytes_safely(
            data
        )


        if len(
            text.strip()
        ) > 20:

            return text


        return ""


    except Exception:

        return ""


# ============================================================
# FILE TYPE DETECTION
# ============================================================

def get_file_extension(
    file_path: str,
) -> str:

    return (

        Path(
            file_path
        )

        .suffix

        .lower()

        .strip()
    )


# ============================================================
# MAIN FILE EXTRACTION ENGINE
# ============================================================

def extract_resume_text(
    file_path: str,
) -> str:

    if not file_path:

        raise ValueError(
            "No resume file path was provided."
        )


    if not os.path.exists(
        file_path
    ):

        raise FileNotFoundError(

            f"Resume file not found: {file_path}"
        )


    extension = get_file_extension(
        file_path
    )


    extraction_methods = {


        ".pdf":

        extract_pdf_text,


        ".docx":

        extract_docx_text,


        ".doc":

        extract_doc_text,


        ".txt":

        extract_txt_text,


        ".rtf":

        extract_rtf_text,


        ".html":

        extract_html_text,


        ".htm":

        extract_html_text,


        ".json":

        extract_json_text,
    }


    extractor = (
        extraction_methods.get(
            extension
        )
    )


    text = ""


    if extractor:

        try:

            text = extractor(
                file_path
            )

        except Exception:

            text = ""


    # Fallback system

    if not text.strip():

        text = extract_generic_text(
            file_path
        )


    text = clean_extracted_text(
        text
    )


    if not text.strip():

        raise ValueError(

            "Unable to extract readable text from this file. "
            "The file may be image-only, encrypted, corrupted, "
            "or use an unsupported proprietary format."
        )


    return text


# ============================================================
# TEXT NORMALIZATION FOR SKILL DETECTION
# ============================================================

def create_searchable_text(
    text: str,
) -> str:

    if not text:

        return ""


    normalized = text.lower()


    # Normalize common separators

    normalized = re.sub(

        r"[_/|,:;()\[\]{}]",

        " ",

        normalized,
    )


    # Normalize dots where appropriate

    normalized = normalized.replace(

        ".",

        " "
    )


    # Collapse whitespace

    normalized = re.sub(

        r"\s+",

        " ",

        normalized,
    )


    return normalized.strip()


# ============================================================
# SKILL VARIATIONS
# ============================================================

SKILL_VARIATIONS = {


    "Python":

    [

        "python",

        "python3",

        "python programming",
    ],


    "JavaScript":

    [

        "javascript",

        "js",

        "ecmascript",
    ],


    "TypeScript":

    [

        "typescript",

        "ts",
    ],


    "React":

    [

        "react",

        "reactjs",

        "react js",
        "react.js",
    ],


    "Node.js":

    [

        "node",

        "nodejs",

        "node js",

        "node.js",
    ],


    "Next.js":

    [

        "nextjs",

        "next js",

        "next.js",
    ],


    "Vue.js":

    [

        "vue",

        "vuejs",

        "vue js",

        "vue.js",
    ],


    "Machine Learning":

    [

        "machine learning",

        "ml",
    ],


    "Artificial Intelligence":

    [

        "artificial intelligence",

        " ai ",
    ],


    "Deep Learning":

    [

        "deep learning",

        "dl",
    ],


    "Natural Language Processing":

    [

        "natural language processing",

        "nlp",
    ],


    "Computer Vision":

    [

        "computer vision",

        "cv",
    ],


    "Large Language Models":

    [

        "large language models",

        "large language model",

        "llm",

        "llms",
    ],


    "Generative AI":

    [

        "generative ai",

        "gen ai",

        "genai",
    ],


    "Object Oriented Programming":

    [

        "object oriented programming",

        "object-oriented programming",

        "oop",
    ],


    "REST APIs":

    [

        "rest api",

        "rest apis",

        "restful api",

        "restful apis",
    ],


    "Tailwind CSS":

    [

        "tailwind",

        "tailwindcss",

        "tailwind css",
    ],


    "Scikit-learn":

    [

        "scikit learn",

        "scikit-learn",

        "sklearn",
    ],


    "Hugging Face":

    [

        "huggingface",

        "hugging face",
    ],


    "Google Cloud":

    [

        "google cloud",

        "gcp",

        "google cloud platform",
    ],


    "AWS":

    [

        "aws",

        "amazon web services",
    ],


    "Microsoft Azure":

    [

        "azure",

        "microsoft azure",
    ],


    "C++":

    [

        "c++",

        "cpp",

        "c plus plus",
    ],


    "C#":

    [

        "c#",

        "csharp",

        "c sharp",
    ],


    "C":

    [

        " c programming ",

        " c language ",

        "programming in c",
    ],
}


# ============================================================
# DETECT SKILLS
# ============================================================

def detect_skills_from_text(
    text: str,
) -> list[str]:

    if not text:

        return []


    searchable_text = (
        " "
        +
        create_searchable_text(
            text
        )
        +
        " "
    )


    detected_skills = []


    # First check variations

    for canonical_skill, variations in (

        SKILL_VARIATIONS.items()

    ):

        for variation in variations:

            variation_normalized = (
                create_searchable_text(
                    variation
                )
            )


            pattern = (

                r"(?<!\w)"

                +

                re.escape(
                    variation_normalized
                )

                +

                r"(?!\w)"
            )


            if re.search(

                pattern,

                searchable_text,

                re.IGNORECASE,
            ):

                detected_skills.append(
                    canonical_skill
                )

                break


    # Check all known skills

    for skill in KNOWN_SKILLS:

        if skill in detected_skills:

            continue


        normalized_skill = (
            create_searchable_text(
                skill
            )
        )


        if not normalized_skill:

            continue


        pattern = (

            r"(?<!\w)"

            +

            re.escape(
                normalized_skill
            )

            +

            r"(?!\w)"
        )


        if re.search(

            pattern,

            searchable_text,

            re.IGNORECASE,
        ):

            detected_skills.append(
                skill
            )


    # Remove duplicates while preserving order

    unique_skills = []


    seen = set()


    for skill in detected_skills:

        key = skill.lower()


        if key not in seen:

            seen.add(
                key
            )

            unique_skills.append(
                skill
            )


    return unique_skills


# ============================================================
# EDUCATION EXTRACTION
# ============================================================

def extract_education(
    text: str,
) -> list[str]:

    education_keywords = [

        "bachelor",

        "master",

        "b.tech",

        "btech",

        "b.e",

        "be ",

        "m.tech",

        "mtech",

        "b.sc",

        "bsc",

        "m.sc",

        "msc",

        "computer science",

        "engineering",

        "university",

        "college",

        "school",
    ]


    lines = text.splitlines()


    education = []


    for line in lines:

        line_clean = line.strip()


        if not line_clean:

            continue


        line_lower = line_clean.lower()


        if any(

            keyword in line_lower

            for keyword in education_keywords

        ):

            education.append(
                line_clean
            )


    return education[:10]


# ============================================================
# EXPERIENCE EXTRACTION
# ============================================================

def extract_experience_signals(
    text: str,
) -> list[str]:

    keywords = [

        "intern",

        "internship",

        "experience",

        "worked",

        "developer",

        "engineer",

        "research",

        "freelance",

        "project",
    ]


    lines = text.splitlines()


    results = []


    for line in lines:

        line_clean = line.strip()


        if not line_clean:

            continue


        line_lower = line_clean.lower()


        if any(

            keyword in line_lower

            for keyword in keywords

        ):

            results.append(
                line_clean
            )


    return results[:15]


# ============================================================
# PROJECT EXTRACTION
# ============================================================

def extract_projects(
    text: str,
) -> list[str]:

    lines = [

        line.strip()

        for line in text.splitlines()

        if line.strip()
    ]


    project_keywords = [

        "project",

        "built",

        "developed",

        "created",

        "designed",

        "implemented",
    ]


    projects = []


    for line in lines:

        line_lower = line.lower()


        if any(

            keyword in line_lower

            for keyword in project_keywords

        ):

            if len(line) >= 10:

                projects.append(
                    line
                )


    return projects[:10]


# ============================================================
# EMAIL EXTRACTION
# ============================================================

def extract_email(
    text: str,
) -> str | None:

    match = re.search(

        r"[A-Za-z0-9._%+-]+@[A-Za-z0-9.-]+\.[A-Za-z]{2,}",

        text,
    )


    if match:

        return match.group(0)


    return None


# ============================================================
# PHONE EXTRACTION
# ============================================================

def extract_phone(
    text: str,
) -> str | None:

    matches = re.findall(

        r"(?:\+?\d{1,3}[-.\s]?)?(?:\(?\d{2,5}\)?[-.\s]?)?\d{5}[-.\s]?\d{4,6}",

        text,
    )


    if matches:

        return matches[0].strip()


    return None


# ============================================================
# MAIN RESUME PARSER
# ============================================================

def parse_resume_text(
    text: str,
) -> dict[str, Any]:

    if not text:

        return {

            "skills": [],

            "education": [],

            "experience_signals": [],

            "projects": [],

            "email": None,

            "phone": None,

            "raw_text": "",

            "text_length": 0,

            "parsing_status":
            "No text provided",
        }


    cleaned_text = clean_extracted_text(
        text
    )


    detected_skills = (
        detect_skills_from_text(
            cleaned_text
        )
    )


    try:

        normalized_skills = (

            normalize_skills(
                detected_skills
            )
        )

    except Exception:

        normalized_skills = (
            detected_skills
        )


    # Final duplicate protection

    final_skills = []


    seen_skills = set()


    for skill in normalized_skills:

        skill_text = str(
            skill
        ).strip()


        if not skill_text:

            continue


        skill_key = (
            skill_text.lower()
        )


        if skill_key not in seen_skills:

            seen_skills.add(
                skill_key
            )

            final_skills.append(
                skill_text
            )


    return {

        "skills":
        final_skills,

        "education":
        extract_education(
            cleaned_text
        ),

        "experience_signals":
        extract_experience_signals(
            cleaned_text
        ),

        "projects":
        extract_projects(
            cleaned_text
        ),

        "email":
        extract_email(
            cleaned_text
        ),

        "phone":
        extract_phone(
            cleaned_text
        ),

        "raw_text":
        cleaned_text,

        "text_length":
        len(
            cleaned_text
        ),

        "skills_detected":
        len(
            final_skills
        ),

        "parsing_status":
        "Resume parsed successfully",
    }
